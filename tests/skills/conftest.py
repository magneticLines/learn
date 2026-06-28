# tests/skills/conftest.py
import pytest
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


@pytest.fixture
def sample_docx(tmp_path):
    """构造一个含命名标题 + Normal 正文的 docx，返回路径。"""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    h1 = doc.add_paragraph("一级标题", style="Heading 1")
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h1.runs:
        r.font.size = Pt(16)
        r.font.bold = True

    body = doc.add_paragraph("这是正文段落。")
    body.paragraph_format.first_line_indent = Pt(24)

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return str(path)
