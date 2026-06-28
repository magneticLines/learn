# tests/skills/test_docx_style.py
import importlib.util
from pathlib import Path
from docx import Document

SCRIPT = Path(".claude/skills/paper-format-check/scripts/docx_style.py")


def _load():
    spec = importlib.util.spec_from_file_location("docx_style", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def test_effective_format_of_body_paragraph(sample_docx):
    ds = _load()
    doc = Document(sample_docx)
    body = [p for p in doc.paragraphs if p.text.startswith("这是正文")][0]
    fmt = ds.effective_format(body)
    assert fmt["font_ascii"] == "Times New Roman"   # 继承自 Normal
    assert fmt["size_pt"] == 12.0
    assert fmt["first_line_indent_pt"] == 24.0


def test_run_override_beats_style(sample_docx):
    ds = _load()
    doc = Document(sample_docx)
    h1 = [p for p in doc.paragraphs if p.text == "一级标题"][0]
    fmt = ds.effective_format(h1)
    assert fmt["size_pt"] == 16.0      # run 覆盖样式
    assert fmt["bold"] is True
    assert fmt["alignment"] == "CENTER"
