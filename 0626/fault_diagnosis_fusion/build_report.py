# -*- coding: utf-8 -*-
"""生成课程大作业报告 docx（题目 8：工业设备故障诊断多传感器融合）。

策略：以模板「大作业报告模版.docx」为基底，保留其封面（含校徽图片）、字体与版式，
仅改写封面题目文字；删除范文正文后，按模板字体字号（章标题黑体15pt居中、二级黑体14pt、
三级黑体12pt、正文宋体小四、图题居中）追加本课题的详细内容，并嵌入 figures/ 下图表。

运行：python build_report.py
输出：../大作业报告_工业故障诊断融合.docx
"""
import os
import json
import shutil

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(ROOT, "figures")
TEMPLATE = os.path.abspath(os.path.join(ROOT, "..", "大作业报告模版.docx"))
OUT = os.path.abspath(os.path.join(ROOT, "..", "大作业报告_工业故障诊断融合.docx"))
M = json.load(open(os.path.join(ROOT, "results", "metrics.json"), encoding="utf-8"))
ACC = M["accuracy"]
FUSED = ACC["三传感器融合"]
BEST = max(ACC["仅振动"], ACC["仅声学"], ACC["仅温度"])
GAIN = (FUSED - BEST) * 100

SONG = "宋体"
HEI = "黑体"
EN = "Times New Roman"


# ---------------- 字体辅助 ----------------
def _set_run(run, cn, size, bold, en=EN, color=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), en); rf.set(qn("w:hAnsi"), en); rf.set(qn("w:eastAsia"), cn)


def _para(doc, text, cn=SONG, size=12, bold=False, align=None, indent_chars=0,
          before=0, after=0, line=1.5, en=EN, page_break_before=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line == 1.5 else None
    if line != 1.5:
        pf.line_spacing = line
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.page_break_before = page_break_before
    if align is not None:
        p.alignment = align
    if indent_chars:
        pf.first_line_indent = Pt(size * indent_chars)
    run = p.add_run(text)
    _set_run(run, cn, size, bold, en)
    return p


def _set_outline(p, level):
    """设置段落大纲级别（0=1级…），使其能被 TOC 域收录。"""
    ppr = p._p.get_or_add_pPr()
    o = ppr.find(qn("w:outlineLvl"))
    if o is None:
        o = OxmlElement("w:outlineLvl"); ppr.append(o)
    o.set(qn("w:val"), str(level))


def chapter(doc, text):
    p = _para(doc, text, cn=HEI, size=15, bold=False,
              align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=8, page_break_before=True)
    _set_outline(p, 0); return p


def sec2(doc, text):
    p = _para(doc, text, cn=HEI, size=14, bold=False, before=10, after=4)
    _set_outline(p, 1); return p


def sec3(doc, text):
    p = _para(doc, text, cn=HEI, size=12, bold=False, before=8, after=2)
    _set_outline(p, 2); return p


def body(doc, text):
    return _para(doc, text, cn=SONG, size=12, indent_chars=2,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_figure(doc, path, caption, width_in=5.6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    _set_run(cap.add_run(caption), SONG, 10.5, False)


# ---------------- 封面编辑 ----------------
def set_para_text(p, text, keep_first_run=True):
    """保留首个 run 的格式，替换段落文本，删除其余 run。"""
    runs = p.runs
    if not runs:
        p.add_run(text); return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def delete_paragraph(p):
    p._element.getparent().remove(p._element)


# ---------------- 主流程 ----------------
def build():
    shutil.copyfile(TEMPLATE, OUT)
    doc = Document(OUT)
    paras = doc.paragraphs

    # 1) 改封面题目与日期（按文本定位，保留字体）
    for p in paras:
        t = p.text.strip().replace("\xa0", " ")
        if t.startswith("题目"):
            set_para_text(p, "题目：  基于振动、声学与红外温度的")
        elif t == "多传感器融合设计":
            set_para_text(p, "工业设备故障多传感器融合诊断设计")
        elif t.startswith("2026") and "月" in t:
            set_para_text(p, "2026 年 6 月")

    # 1b) 改页眉：模板页眉为范文标题“……行人重识别”，改为本课题标题；
    #     首页（封面）不显示页眉。
    HEADER_TEXT = "基于振动、声学与红外温度的工业设备故障多传感器融合诊断"
    for si, sec in enumerate(doc.sections):
        if si == 0:
            # 封面（首页）不显示页眉：启用首页独立页眉并清空它
            sec.different_first_page_header_footer = True
            fph = sec.first_page_header
            fph.is_linked_to_previous = False
            for hp in fph.paragraphs:
                set_para_text(hp, "")
        for hdr in (sec.header, sec.even_page_header):       # 正文页眉改为本课题标题
            for hp in hdr.paragraphs:
                if hp.text.strip() and ("行人" in hp.text or "transformer" in hp.text.lower()
                                        or "重识别" in hp.text):
                    set_para_text(hp, HEADER_TEXT)

    # 2) 删除范文正文：从“摘要”所在段落到文档末尾全部删除
    cut = None
    for i, p in enumerate(paras):
        if p.text.strip() == "摘要":
            cut = i; break
    if cut is None:                       # 兜底：找第一个含“摘要”的段
        for i, p in enumerate(paras):
            if "摘要" in p.text:
                cut = i; break
    for p in paras[cut:]:
        delete_paragraph(p)

    # 2b) 删除模板自带的自动目录（包在 w:sdt 内容控件里，不在 doc.paragraphs 中，
    #     上面的按段删除删不到它，会残留为重复目录）。仅删含 TOC 域的 sdt。
    body_el = doc.element.body
    for sdt in body_el.findall(".//" + qn("w:sdt")):
        instr = sdt.findall(".//" + qn("w:instrText"))
        if any("TOC" in (e.text or "") for e in instr):
            sdt.getparent().remove(sdt)

    # 2c) 清理封面尾部残留的空段落与分页符，避免出现空白页
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        delete_paragraph(doc.paragraphs[-1])
    if doc.paragraphs:                       # 去掉最后一段里可能残留的分页符
        for br in doc.paragraphs[-1]._p.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)

    # 3) 追加本课题内容（各章节用“段前分页”起新页，避免多余空白页）
    _build_abstract(doc)
    _build_toc(doc)
    _build_ch1(doc)
    _build_ch2(doc)
    _build_ch3(doc)
    _build_ch4(doc)
    _build_ch5(doc)
    _build_ch6(doc)
    _build_refs(doc)

    doc.save(OUT)
    print("报告已生成:", OUT)


def chapter_plain(doc, text):
    """章级标题样式但不进目录（用于摘要、目录页自身标题）。"""
    return _para(doc, text, cn=HEI, size=15, bold=False,
                 align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=8, page_break_before=True)


def _build_abstract(doc):
    chapter_plain(doc, "摘  要")
    body(doc,
         "工业旋转设备的滚动轴承长期运行于高速、重载、变工况环境，是设备故障的高发部位，"
         "其故障若不能及时、准确地诊断，轻则造成非计划停机，重则引发安全事故。传统故障诊断"
         "多依赖单一振动传感器，受信息维度有限、易受工况与噪声干扰等因素制约，难以在复杂"
         "场景下同时兼顾诊断准确率与鲁棒性。针对该问题，本文设计并实现了一套基于振动、声学"
         "与红外温度三类异构传感器的特征级融合故障诊断方法。")
    body(doc,
         "受实验条件限制，本文依据滚动轴承故障机理对三类传感器信号进行建模与模拟生成，"
         "覆盖正常、内圈故障、外圈故障、滚动体故障四种典型工况，每类 200 个样本。系统分别"
         "从振动信号提取时域（均方根、峭度、峰值因子、标准差）与频域（轴承故障特征频率"
         " BPFI/BPFO/BSF 频带能量、频谱重心）特征，从声学信号提取均方根、频谱熵、高频能量"
         "占比与异响频带能量特征，从红外温度提取均值、最大值、最大温升梯度、总温升与热点数"
         "特征，共计 18 维。采用特征级融合策略将三类传感器特征拼接为统一特征向量，并以随机"
         "森林作为分类器实现四工况诊断。")
    body(doc,
         "对比实验表明：仅振动、仅声学、仅温度的故障识别准确率分别为 "
         f"{ACC['仅振动']:.1%}、{ACC['仅声学']:.1%}、{ACC['仅温度']:.1%}，"
         f"而三传感器融合准确率达 {FUSED:.1%}，较最佳单传感器提升约 {GAIN:.1f} 个百分点。"
         "混淆矩阵与特征重要性分析进一步表明，三类传感器在“故障类型—故障有无—故障严重程度”"
         "三个维度上信息互补，融合后各类别误判显著减少。实验结果验证了多传感器融合在提升"
         "故障诊断准确率与鲁棒性方面的有效性。")
    p = doc.add_paragraph(); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(6)
    _set_run(p.add_run("关键词："), HEI, 12, False)
    _set_run(p.add_run("多传感器融合；故障诊断；滚动轴承；特征级融合；随机森林"), SONG, 12, False)


def _build_toc(doc):
    chapter_plain(doc, "目  录")
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "（请在 Word 中右键“更新域”生成目录）"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for e in (b, instr, sep, t, end):
        run._r.append(e)
    _set_run(run, SONG, 12, False)


def _build_ch1(doc):
    chapter(doc, "第1章 绪论")
    sec2(doc, "1.1 研究背景与意义")
    body(doc, "随着工业自动化与智能制造的快速发展，旋转机械在电力、冶金、石化、轨道交通与"
              "高端装备制造等领域得到广泛应用。滚动轴承作为旋转机械中传递载荷与支撑转动的"
              "核心部件，其运行状态直接关系到整机的安全性与可靠性。据统计，旋转机械约 30% 以上"
              "的故障源于轴承失效。轴承长期工作于高速、重载、冲击与变工况环境，疲劳剥落、磨损、"
              "点蚀等损伤难以避免，一旦发生故障而未被及时发现，轻则导致设备非计划停机、生产中断，"
              "重则引发连锁损坏甚至安全事故，造成重大经济损失。")
    body(doc, "因此，对旋转设备开展状态监测与故障诊断，实现故障的早期发现与准确定位，对于保障"
              "设备安全运行、降低维护成本、推动由“定期维修”向“预测性维护”转变具有重要的工程"
              "价值与现实意义。")
    sec3(doc, "1.1.1 单一传感器诊断的局限")
    body(doc, "传统故障诊断方法多依赖单一振动传感器。然而，单一信息源所能获取的状态信息维度"
              "有限：振动信号对故障类型敏感、可通过特征频率区分故障部位，但易受机械噪声、负载"
              "波动与传感器安装方式的干扰；声学信号便于非接触测量、对早期异响敏感，但其传播易"
              "受环境噪声污染，且不同故障类型的声学特征区分度较低；红外温度能够反映设备的整体"
              "热状态与润滑情况，但对早期局部故障响应滞后。单一传感器难以在复杂工况下同时保证"
              "诊断的准确性与鲁棒性。")
    sec3(doc, "1.1.2 多传感器融合的价值")
    body(doc, "多传感器融合通过综合利用多个传感器在时间或空间上的冗余与互补信息，可获得比"
              "单一传感器更全面、更可靠的状态估计。将振动、声学与温度三类异构信息有机结合，"
              "有望在保留各自优势的同时弥补彼此短板，从而突破单一传感器的性能瓶颈，提升故障"
              "诊断的准确率与抗干扰能力。这正是本文研究的出发点。")
    sec2(doc, "1.2 国内外研究现状")
    sec3(doc, "1.2.1 旋转机械故障诊断方法")
    body(doc, "故障诊断方法大体经历了从信号处理到智能诊断的发展。基于信号处理的方法包括时域"
              "统计分析、频谱分析、包络解调、小波变换与经验模态分解等，能够有效提取轴承故障"
              "特征频率；近年来，以支持向量机、随机森林为代表的机器学习方法以及以卷积神经网络、"
              "循环神经网络为代表的深度学习方法显著提升了诊断的自动化程度与精度，但通常需要"
              "较多的标注样本与计算资源。")
    sec3(doc, "1.2.2 多传感器信息融合")
    body(doc, "按融合发生的层次，多传感器融合可分为数据级、特征级与决策级三类。数据级融合直接"
              "融合原始信号，信息保留最完整但对时空配准要求高、计算量大；决策级融合融合各传感器"
              "的独立判决结果，实现简单、容错性好，但损失了细节信息；特征级融合融合各传感器提取"
              "的特征向量，在信息保留与实现复杂度之间取得较好平衡，是工程实践中应用最为广泛的"
              "方案。")
    sec3(doc, "1.2.3 研究现状总结")
    body(doc, "总体来看，现有研究多聚焦于单一信号或两类信号的融合，针对振动—声学—红外温度三类"
              "异构传感器在统一框架下进行特征级融合诊断、并系统量化对比融合增益的工作仍相对"
              "有限。本文即面向这一问题展开设计与验证。")
    sec2(doc, "1.3 本文主要工作")
    body(doc, "本文主要工作包括：①依据滚动轴承故障机理，建模并模拟生成振动、声学、红外温度三类"
              "传感器在四种典型工况下的信号数据；②针对三类异构信号分别设计时域、频域与统计特征"
              "提取方法，构建 18 维特征体系；③构建特征级融合框架，以随机森林实现四工况故障分类；"
              "④通过单传感器基线与融合的系统对比实验，从准确率、混淆矩阵与特征重要性三个角度"
              "量化评估融合的有效性；⑤实现可交互的实时诊断演示系统，便于结果展示。")


def _build_ch2(doc):
    chapter(doc, "第2章 相关理论基础")
    sec2(doc, "2.1 多传感器融合层次")
    body(doc, "多传感器融合是指综合利用多个传感器获取的信息，通过一定的准则进行关联、估计与"
              "组合，以获得对被监测对象更准确、更可靠描述的过程。按融合层次可分为数据级、特征级"
              "与决策级融合。本文采用特征级融合：先对各传感器信号独立提取特征，再将特征向量按"
              "维度拼接为统一的融合特征，最后送入分类器进行决策。该层次既保留了较丰富的细节"
              "信息，又避免了数据级融合对原始信号严格时空配准的依赖，工程上易于实现。")
    sec2(doc, "2.2 滚动轴承故障特征频率")
    body(doc, "当滚动轴承的内圈、外圈或滚动体表面出现局部损伤时，滚动体周期性地碰撞损伤点，"
              "会在振动信号中激起以特定频率出现的周期性冲击，称为故障特征频率。其主要包括内圈"
              "故障频率 BPFI、外圈故障频率 BPFO 与滚动体故障频率 BSF，三者由轴承几何参数（节圆"
              "直径、滚动体直径、接触角、滚动体个数）与转速共同决定。由于不同部位的故障特征频率"
              "不同，故障特征频率是区分故障类型（内圈/外圈/滚动体）的关键依据，构成本文振动频域"
              "特征的物理基础。")
    sec2(doc, "2.3 特征提取相关指标")
    body(doc, "时域统计指标方面，均方根反映信号能量水平；峭度对冲击成分敏感，常用于早期故障"
              "检测；峰值因子刻画信号峰值相对能量的突出程度。频域方面，故障特征频带能量占比"
              "反映对应故障的严重程度，频谱重心刻画能量在频率轴上的分布。声学方面，频谱熵刻画"
              "信号频率成分的无序程度，高频能量占比对异响敏感。温度方面，温升梯度与热点数能够"
              "反映摩擦加剧导致的异常发热。")
    sec2(doc, "2.4 随机森林分类器")
    body(doc, "随机森林是一种基于决策树的集成学习算法。其通过对训练样本进行有放回抽样（Bagging）"
              "并在每个节点随机选取特征子集来训练多棵相互独立的决策树，最终以多数投票方式输出"
              "分类结果。随机森林具有抗过拟合能力强、对特征量纲不敏感、无需复杂归一化、可输出"
              "特征重要性等优点，特别适合处理本文这类由多源异构传感器构成的中等维度特征的分类"
              "问题，因而被选作本文的融合分类器。")


def _build_ch3(doc):
    chapter(doc, "第3章 系统设计与数据生成")
    sec2(doc, "3.1 系统总体框架")
    body(doc, "本文设计的多传感器融合故障诊断系统总体流程为：多传感器数据采集（模拟生成）→ "
              "各传感器特征提取 → 特征级融合 → 随机森林分类 → 诊断结果输出与可视化。三类传感器"
              "并行采集，经各自的特征提取后拼接为统一特征向量，由随机森林输出四种工况之一的"
              "诊断结果。系统按模块化设计，各模块职责单一、接口清晰，便于维护与扩展。")
    sec2(doc, "3.2 多传感器数据生成")
    body(doc, "在缺乏真实传感器与实测数据的条件下，为支撑方法验证，本文依据滚动轴承故障机理"
              "对三类传感器信号进行建模与模拟生成。四种工况（正常、内圈故障、外圈故障、滚动体"
              "故障）各生成 200 个样本，共 800 个样本；振动与声学信号采样率取 12 kHz、每样本"
              "2048 点，温度为长度 60 的时序。")
    sec3(doc, "3.2.1 振动信号建模")
    body(doc, "振动信号由三部分叠加而成：①转频及其谐波，模拟设备正常运转的固有振动；②故障"
              "特征频率处的周期性冲击与包络谱线，模拟轴承不同部位损伤激起的冲击响应，其特征"
              "频率随故障类型（BPFI/BPFO/BSF）而不同；③高斯测量噪声。为贴近真实并体现难度，"
              "故障冲击幅值在各类型间相近且带随机扰动，使弱故障样本可能与正常样本混淆。")
    sec3(doc, "3.2.2 声学与温度信号建模")
    body(doc, "声学信号在振动声辐射（取振动信号的高频强调分量）的基础上，对故障工况叠加固定"
              "频带的异响分量与宽带突发噪声，并叠加较强的环境噪声；由于三种故障的异响形态相近，"
              "声学信号主要可区分“有无故障”而难以分辨故障类型。温度信号按一阶惯性环节建模其"
              "升温到稳态的过程，对故障工况设置更高的稳态温度、更快的温升速率并叠加随机热点"
              "尖峰，但三种故障的稳态温度接近，故温度主要反映故障的有无与大致严重程度。")
    add_figure(doc, os.path.join(FIG, "fig1_vibration_waveform.png"),
               "图 3-1  四种工况的振动时域波形")
    add_figure(doc, os.path.join(FIG, "fig2_vibration_spectrum.png"),
               "图 3-2  振动频谱与故障特征频率（BPFI/BPFO/BSF）")
    add_figure(doc, os.path.join(FIG, "fig3_acoustic_temperature.png"),
               "图 3-3  各工况声学信号与红外温度曲线")
    sec2(doc, "3.3 特征提取")
    body(doc, "针对三类信号分别设计特征提取方法，合计 18 维。振动特征 9 维：时域的均方根、峭度、"
              "峰值因子、标准差，频域的 BPFI/BPFO/BSF 频带能量占比与频谱重心；声学特征 4 维："
              "均方根、频谱熵、高频能量占比、异响频带能量；温度特征 5 维：均值、最大值、最大"
              "温升梯度、总温升、热点数。三类特征在工况区分能力上各有侧重，分别对应“故障类型”"
              "“故障有无”与“故障严重程度”，构成互补关系。")
    sec2(doc, "3.4 特征级融合与分类")
    body(doc, "将振动、声学、温度三类特征向量按维度拼接，形成 18 维融合特征向量，送入随机森林"
              "（200 棵决策树）进行四分类。为评估融合效果，另以各单类传感器的特征子集分别训练"
              "结构相同的随机森林作为基线，从而在同一实验条件下对比单传感器与融合的诊断性能。")


def _build_ch4(doc):
    chapter(doc, "第4章 实验与结果分析")
    sec2(doc, "4.1 实验设置")
    body(doc, f"数据集共 {M['n_samples']} 个样本（4 类 × 200），按 7∶3 的比例划分为训练集"
              f"（{M['split']['train']} 个）与测试集（{M['split']['test']} 个），并采用分层抽样"
              "保证各类别比例均衡。分类器为随机森林（200 棵树，随机种子固定以保证可复现）。"
              "评价指标采用故障识别准确率，并辅以混淆矩阵分析各类别的识别情况、以特征重要性"
              "分析各特征的贡献。")
    sec2(doc, "4.2 单传感器与融合准确率对比")
    body(doc, "四种配置在测试集上的故障识别准确率如表 4-1 与图 4-1 所示。可以看出，仅振动配置"
              f"准确率最高（{ACC['仅振动']:.1%}），因其能够借助故障特征频率较好地区分故障类型；"
              f"仅声学与仅温度配置准确率较低（分别为 {ACC['仅声学']:.1%}、{ACC['仅温度']:.1%}），"
              "主要因为二者分别只能判断“有无故障”与“严重程度”，难以分辨具体的故障类型。"
              f"三传感器融合充分利用了三类信息的互补性，准确率达到 {FUSED:.1%}，优于任一单"
              f"传感器，较最佳单传感器提升约 {GAIN:.1f} 个百分点，验证了多传感器融合的有效性。")
    _add_result_table(doc)
    add_figure(doc, os.path.join(FIG, "fig4_accuracy_comparison.png"),
               "图 4-1  单传感器与多传感器融合准确率对比", width_in=4.6)
    sec2(doc, "4.3 混淆矩阵分析")
    body(doc, "图 4-2 给出四种配置在测试集上的混淆矩阵。可以看到，仅声学、仅温度配置将三种"
              "故障类型大量相互混淆，对角线占比偏低；仅振动配置的故障类型区分较好，主要的误判"
              "集中在正常与轻微故障之间。经过三传感器融合后，各类别的对角线占比显著提高，类型"
              "间与正常—故障间的误判均明显减少，表明融合有效整合了不同传感器的判别能力。")
    add_figure(doc, os.path.join(FIG, "fig5_confusion_matrices.png"),
               "图 4-2  各配置混淆矩阵", width_in=5.2)
    sec2(doc, "4.4 特征重要性分析")
    body(doc, "图 4-3 给出融合模型输出的特征重要性排序。结果显示，振动频域的故障特征频带能量、"
              "声学的异响频带能量与高频能量占比、温度的温升与热点相关特征均位居前列，说明三类"
              "传感器的特征对最终诊断均有实质性贡献。这从特征层面印证了三类传感器信息互补、"
              "多传感器融合合理且必要的结论。")
    add_figure(doc, os.path.join(FIG, "fig6_feature_importance.png"),
               "图 4-3  融合模型 Top-10 特征重要性", width_in=4.8)


def _build_ch5(doc):
    chapter(doc, "第5章 系统实现")
    sec2(doc, "5.1 软件总体结构")
    body(doc, "系统采用 Python 实现，依赖 NumPy、SciPy、scikit-learn 与 Matplotlib 等开源库，"
              "按职责划分为五个模块：数据生成模块（mock_sensors）负责按机理模拟三类传感器信号；"
              "特征提取模块（features）负责时域、频域与温度特征的提取；融合与对比实验模块"
              "（fusion）负责单传感器基线与特征级融合模型的训练与评估；可视化模块（visualize）"
              "负责生成各类结果图表；端到端主流程（pipeline）将上述环节串联，一键完成从数据"
              "生成到结果输出的全过程。")
    sec2(doc, "5.2 核心模块实现")
    sec3(doc, "5.2.1 数据生成与特征提取")
    body(doc, "数据生成模块以可配置的工况参数（故障特征频率、冲击幅值、稳态温度等）为输入，"
              "结合随机扰动生成具有合理重叠的样本；特征提取模块对每个样本依次计算 18 维特征，"
              "并标注各特征所属的传感器，以支持单传感器子集实验。")
    sec3(doc, "5.2.2 融合分类与评估")
    body(doc, "融合模块按传感器归属选取特征列，分别构建“仅振动/仅声学/仅温度/三传感器融合”四种"
              "配置，使用相同结构的随机森林进行训练，并统一在测试集上计算准确率、混淆矩阵与"
              "分类报告，保证对比的公平性。")
    sec2(doc, "5.3 实时诊断演示")
    body(doc, "为便于结果展示与答辩演示，系统实现了交互式诊断界面：点击按钮即可随机抽取一个"
              "工况样本，以动画方式逐步绘制振动、声学、温度三路信号，并实时给出融合模型的诊断"
              "结果，诊断正确与否以颜色加以区分。该演示可用于录制演示视频，直观展示系统的工作"
              "过程与诊断效果。")
    sec2(doc, "5.4 功能测试")
    body(doc, "对数据生成、特征提取、模型训练与诊断流程进行了端到端测试：主流程可一键稳定运行"
              "并产出全部图表与量化指标；交互式演示可正常抽样与诊断。测试结果表明系统功能完整、"
              "运行稳定、结果可复现。")


def _build_ch6(doc):
    chapter(doc, "第6章 总结与展望")
    sec2(doc, "6.1 工作总结")
    body(doc, "本文面向工业旋转设备的滚动轴承故障诊断，设计并实现了基于振动、声学与红外温度"
              "三类异构传感器的特征级融合诊断系统，完成了从数据机理建模、特征提取、特征级融合"
              "到随机森林分类的完整流程，并通过系统的对比实验验证了方法的有效性。实验表明，"
              f"三传感器融合的故障识别准确率达 {FUSED:.1%}，较最佳单传感器提升约 {GAIN:.1f} "
              "个百分点。")
    sec2(doc, "6.2 主要创新点")
    sec3(doc, "6.2.1 统一的三传感器特征级融合框架")
    body(doc, "构建了振动—声学—红外温度三类异构传感器在统一框架下的特征级融合诊断流程，实现"
              "了多源异构信息的有机整合。")
    sec3(doc, "6.2.2 体现互补性的特征体系设计")
    body(doc, "设计了能够体现三类传感器“故障类型—故障有无—故障严重程度”互补关系的 18 维特征"
              "体系，为融合增益提供了清晰的物理与数据支撑。")
    sec3(doc, "6.2.3 多角度的融合有效性验证")
    body(doc, "从准确率、混淆矩阵与特征重要性三个角度，系统量化验证了多传感器融合相较单传感器"
              "的性能提升，使结论更具说服力。")
    sec2(doc, "6.3 不足与展望")
    body(doc, "本文仍存在以下不足：其一，实验数据为依据机理模拟生成，与真实工况数据在分布上仍"
              "存在差距；其二，融合方式为特征级简单拼接，未充分建模传感器之间的时序关联与可靠度"
              "差异。未来可在以下方向继续深入：①引入 CWRU 等公开真实数据集进行验证与迁移；"
              "②探索决策级融合、基于注意力的加权融合等更先进的融合策略；③引入深度学习方法"
              "自动学习多源特征表示，以进一步提升复杂工况下的诊断性能与泛化能力。")


def _build_refs(doc):
    chapter(doc, "参考文献")
    refs = [
        "[1] 雷亚国, 贾峰, 孔德同, 等. 大数据下机械智能故障诊断的机遇与挑战[J]. 机械工程学报, 2018, 54(5): 94-104.",
        "[2] Randall R B, Antoni J. Rolling element bearing diagnostics—A tutorial[J]. Mechanical Systems and Signal Processing, 2011, 25(2): 485-520.",
        "[3] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "[4] 韩崇昭, 朱洪艳, 段战胜. 多源信息融合[M]. 2版. 北京: 清华大学出版社, 2010.",
        "[5] Smith W A, Randall R B. Rolling element bearing diagnostics using the Case Western Reserve University data: A benchmark study[J]. Mechanical Systems and Signal Processing, 2015, 64-65: 100-131.",
        "[6] 张西宁, 雷威, 余迪. 多传感器信息融合在旋转机械故障诊断中的应用综述[J]. 振动与冲击, 2020, 39(20): 1-12.",
        "[7] Lei Y, Yang B, Jiang X, et al. Applications of machine learning to machine fault diagnosis: A review and roadmap[J]. Mechanical Systems and Signal Processing, 2020, 138: 106587.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        _set_run(p.add_run(r), SONG, 10.5, False)


def _add_result_table(doc):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(cap.add_run("表 4-1  各配置故障识别准确率"), SONG, 10.5, False)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for c, txt in zip(table.rows[0].cells, ["配置", "故障识别准确率"]):
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(c.paragraphs[0].add_run(txt), HEI, 11, False)
    for name in ["仅振动", "仅声学", "仅温度", "三传感器融合"]:
        cells = table.add_row().cells
        bold = name == "三传感器融合"
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cells[0].paragraphs[0].add_run(name), SONG, 11, bold)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cells[1].paragraphs[0].add_run(f"{ACC[name]:.1%}"), SONG, 11, bold)
    doc.add_paragraph()


if __name__ == "__main__":
    build()
